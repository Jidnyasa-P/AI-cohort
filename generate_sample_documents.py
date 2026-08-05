from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from PIL import Image, ImageDraw, ImageFont

root = Path("sample_documents")
root.mkdir(exist_ok=True)

pdf_path = root / "summary_of_benefits_and_coverage.pdf"
docx_path = root / "claims_process.docx"
scan_png_path = root / "enrollment_form_scan.png"
scan_pdf_path = root / "enrollment_form_scan.pdf"

summary_title = "Summary of Benefits and Coverage"
summary_lines = [
    "Plan: Acme Health Premier",
    "Effective Date: 01/01/2026",
    "",
    "Covered Services:",
    " - Primary care visits: $20 copay",
    " - Specialist visits: $40 copay",
    " - Preventive care: 100% covered",
    " - Prescription drugs: Tier 1 $10 copay, Tier 2 $30 copay",
    "",
    "Out-of-Pocket Maximum:",
    " - Individual: $3,000",
    " - Family: $6,000",
    "",
    "Notes:",
    "This document is a summary only. For full coverage details, refer to the plan policy document."
]

c = canvas.Canvas(str(pdf_path), pagesize=letter)
width, height = letter
c.setFont("Helvetica-Bold", 18)
c.drawString(72, height - 72, summary_title)

c.setFont("Helvetica", 11)
y = height - 108
for line in summary_lines:
    c.drawString(72, y, line)
    y -= 18
    if y < 72:
        c.showPage()
        c.setFont("Helvetica", 11)
        y = height - 72
c.save()

# Create Word document
document = Document()
document.add_heading("Claims Process", level=1)
document.add_paragraph(
    "This document describes the steps members should follow when submitting a claim. "
    "It is intended as a general template for internal or member-facing materials."
)
document.add_heading("Step 1: Review Coverage", level=2)
document.add_paragraph(
    "Confirm that the service or treatment is covered under the member's plan before submitting a claim. "
    "Check deductible, copay, and prior authorization requirements."
)
document.add_heading("Step 2: Gather Documentation", level=2)
document.add_paragraph(
    "Collect itemized bills, provider notes, receipts, and any referral/authorization documents. "
    "Ensure all provider information and dates of service are included."
)
document.add_heading("Step 3: Submit the Claim", level=2)
document.add_paragraph(
    "Submit the form and supporting documents through the web portal, fax, or mail. "
    "Include the member ID and claim type on every page."
)
document.add_heading("Step 4: Track the Claim", level=2)
document.add_paragraph(
    "Monitor claim status through the member portal or customer service. "
    "Respond promptly to any requests for additional information."
)
document.save(str(docx_path))

# Create scanned enrollment form image
width, height = 1654, 2339
image = Image.new("RGB", (width, height), "#f6f3eb")
draw = ImageDraw.Draw(image)

try:
    font = ImageFont.truetype("arial.ttf", 28)
    font_small = ImageFont.truetype("arial.ttf", 20)
except Exception:
    font = ImageFont.load_default()
    font_small = ImageFont.load_default()

margin_x = 100
margin_y = 100
line_height = 42

draw.text((margin_x, margin_y), "Enrollment Form", fill="black", font=font)
y = margin_y + 80

fields = [
    "Member Name:",
    "Date of Birth:",
    "Member ID:",
    "Address:",
    "City/State/ZIP:",
    "Phone Number:",
    "Email:",
    "Coverage Start Date:",
    "Plan Selection:",
    "Signature:",
    "Date Signed:"
]

for field in fields:
    draw.text((margin_x, y), field, fill="black", font=font_small)
    y += line_height * 1.8
    draw.line((margin_x, y, width - margin_x, y), fill="#555555", width=2)
    y += line_height * 0.5

notes_y = y + 20
notes_text = (
    "Please print clearly. This scanned form is for illustrative purposes only. "
    "Submit completed forms via secure upload or mail."
)
draw.text((margin_x, notes_y), notes_text, fill="#333333", font=font_small)

image.save(str(scan_png_path), dpi=(200, 200))
image.convert("RGB").save(str(scan_pdf_path), "PDF", resolution=200.0)

print(f"Generated: {pdf_path}")
print(f"Generated: {docx_path}")
print(f"Generated: {scan_png_path}")
print(f"Generated: {scan_pdf_path}")
